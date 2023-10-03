"""
특정 시트의 셀들을 읽어서
docx의 테이블로 만드는 작업을 합니다.
"""

def xlsx_sheet_to_docx_table():
    import docx
    from simple_term_menu import TerminalMenu
    from openpyxl import load_workbook
    from docx.shared import Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from utils import print_summary, create_dir_if_not_exist
    import os

    def init_console():
        base_path = 'services/xlsx_sheet_to_docx_table'
        base_input_path = os.path.abspath(f'{base_path}/xlsx_sheet_to_docx_table_input')
        base_output_path = os.path.abspath(f'{base_path}/xlsx_sheet_to_docx_table_output')
        create_dir_if_not_exist(base_output_path)

        print("변환하려는 XLSX 파일 이름을 입력해 주세요(기본 이름 - 도서대장):", end = ' ')
        input_file_input = input()
        input_file_input = '도서대장' if len(input_file_input.strip()) == 0 else input_file_input.strip()
        input_filename = f'{base_input_path}/{input_file_input}.xlsx'
        print("변환 결과 DOCX 파일 이름을 입력해 주세요:", end = ' ')
        output_filename = f'{base_output_path}/{input().strip()}.docx'

        return (input_filename, output_filename)

    """Load Excel"""
    def parse_excel_data(input_filename, output_filename):
        
        excel_file = load_workbook(input_filename)
        sheetnames = excel_file.sheetnames
        terminal_menu = TerminalMenu(sheetnames, title="작업할 시트를 선택해 주세요.")
        menu_entry_index = terminal_menu.show()

        target_sheet_name = sheetnames[menu_entry_index]
        target_sheet = excel_file[target_sheet_name]
        
        
        data = []
        # (min_row=2, min_col=1, max_row=6, max_col=3)
        for row in target_sheet.iter_rows(min_row=2, min_col=1):
            tmp = []
            count = 0
            for cell in row:
                if cell.value == None: 
                    count += 1
                tmp.append(f'{cell.value}')
            if count == len(row):
                break
            if len(tmp) > 0:
                data.append(tmp)

        # for row in data:
        #     print(row)
        return data

    """Write Document"""
    def save_docx(data):
        doc = docx.Document()
        
        # doc.add_heading('GeeksForGeeks', 0)
        # rows=len(data), cols=len(data[0])
        col_len = len(data[0])
        table = doc.add_table(rows=0, cols=col_len)
        
        widths = (
            Inches(1), 
            Inches(2), 
            Inches(3), 
            Inches(3), 
            Inches(2), 
            Inches(2), 
            Inches(2)
        )
        
        for i in range(len(data)):
            row = table.add_row().cells
            for j in range(col_len):
                row[j].text = data[i][j]
                row[j].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
                row[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            table.rows[i].height = Inches(1)
        
        for i in range(col_len):
            table.columns[i].width = widths[i]
        table.style = 'Table Grid'
        doc.save(output_filename)

    try:
        (input_filename, output_filename) = init_console()
        parsed_data = parse_excel_data(input_filename, output_filename)
        save_docx(parsed_data)

        summary = {
            "변환 대상 파일": input_filename, 
            "변환 결과 파일": output_filename
        }
        print_summary(summary)

    except FileNotFoundError: 
        print("해당 파일을 찾을 수 없습니다.")
    except KeyError:
        print("올바른 파일인지 확인해 주세요.")
    except UnicodeDecodeError as e:
        print(e)

if __name__ == "__main__":
    # import sys
    # input_filename = sys.argv[1]
    # output_filename = sys.argv[2]
    # xlsx_sheet_to_docx_table(f'./pdf_to_docx_input/{input_filename}.xlsx', f'./pdf_to_docx_output/{output_file}.docx')
    xlsx_sheet_to_docx_table()